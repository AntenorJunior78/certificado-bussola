# Importar as bibliotecas necessárias
from tkinter.ttk import*  # Imports all widgets and advanced styles from the ttk (Themed Tkinter) library.
import customtkinter  # Imports the customtkinter module, which may contain customizations or custom widgets.
import time  # Imports the time module, used for pausing execution and creating delay effects.
import pandas as pd  # Imports the pandas module and renames it to pd, used for working with data structures and data analysis.
from io import StringIO  # Imports the StringIO class from the io module, used for working with strings as files.
from tkinter import ttk
from io import BytesIO
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import tkinter as tk
import rpy2.robjects as ro


from pandastable.data import TableModel

from tkinter import *# Imports all modules and classes from the tkinter library, used to create graphical interfaces.

from PIL import Image, ImageTk
import sys  # Adicione esta linha para importar o módulo sys
import time
import numpy as np
from pandastable import Table
from matplotlib.figure import Figure
from matplotlib.backends.backend_tkagg import (FigureCanvasTkAgg, NavigationToolbar2Tk)
from time import strftime
import os
from io import StringIO
from datetime import datetime
from tkinter.filedialog import askdirectory
import tkinter.messagebox
import webbrowser
import platform

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from datetime import datetime

from docx.shared import Cm
import docx 





def run():
    
    root = Tk()
    root.title('Geo-Resistivity-Meter')
    width_of_window = 660
    height_of_window = 600
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()
    x_coordinate = (screen_width/2)-(width_of_window/2)
    y_coordinate = (screen_height/2)-(height_of_window/2)
    root.geometry("%dx%d+%d+%d" %(width_of_window,height_of_window,x_coordinate,y_coordinate))
    root.configure(bg='white')

    icon = PhotoImage(file="figure/energia3.png")#logo
    root.iconphoto(True, icon)
        
    my_font=('arial', 20, 'bold')
    my_font1=('fantasy', 30, 'bold')
    my_font2=('arial', 12, 'bold')
    root.resizable(False,False)
    # ==============================================================================
    # ================== ABAS    ===================================================
    # ==============================================================================
    s = ttk.Style()
    s.theme_create('pastel', settings={
            ".": {
                "configure": {
                    "background": 'white', # All except tabs
                    "font": 'red'
                }
            },
            "TNotebook": {
                "configure": {
                    "background":'#0b8be0', # Your margin color
                    "tabmargins": [2, 5, 0, 0], # margins: left, top, right, separator
                }
            },
            "TNotebook.Tab": {
                "configure": {
                    "background": '#d9ffcc', # tab color when not selected
                    "padding": [10, 2], # [space between text and horizontal tab-button border, space between text and vertical tab_button border]
                    "font":"white"
                },
                "map": {
                    "background": [("selected", '#ccffff')], # Tab color when selected
                    "expand": [("selected", [1, 1, 1, 0])] # text margins
                }
            }
    })
    s.theme_use('pastel')
    s.configure('TNotebook.Tab', font=my_font2 , highlightbackground="white")#fonte da abas
    tabsystem = ttk.Notebook(root)

    # Create new tabs using Frame widget
    tab1 = Frame(tabsystem,width=995,height=800)
    tabsystem.add(tab1, text='  Principal  ')
    tabsystem.place(x=5,y=12) 
    
        
    
    def obter_data():
        dia = entrada_dia.get()
        mes = entrada_mes.get()
        ano = entrada_ano.get()
        data_formatada = f"{dia}/{mes}/{ano}"
        resultado_label.config(text=f"Data inserida: {data_formatada}")
    
    def novaTab():
        
        global tab3
        tab3 = Frame(tabsystem,width=995,height=800)
        tabsystem.add(tab3, text='  PROCEDIMENTOS SEGUIDOS  ')
        tabsystem.place(x=5,y=12)
        Frame(tab3, width=642, height=550, background="#dce6f4", highlightbackground="black", highlightthickness=3).place(x=3, y=3)
        
        
        labelframe10 = LabelFrame(tab3, text=" ETAPAS ",bg='#dce6f4',font=my_font2)
        labelframe10.place(x=10, y=10)
        labelframe11 = LabelFrame(tab3, text=" INSTRUMENTOS E LOCAIS ",bg='#dce6f4',font=my_font2)
        labelframe11.place(x=10, y=190)
        labelframe12 = LabelFrame(tab3, text=" CALIBRAÇÃO ",bg='#dce6f4',font=my_font2)
        labelframe12.place(x=10, y=370)
        
        

        # Criar uma caixa de texto com 10 linhas usando Text
        texto_inicial = (
        "As medições foram realizadas em duas etapas. "
        "A primeira no Laboratório de Desenvolvimento de Sensores Magnéticos (LDSM/ON) com instrumentos acreditados. "
        "A segunda no Observatório Magnético de Vassouras(VSS) em operação desde 1915 e referenciado pela IAGA (International Association of Geomagnetism and Aeronomy). "
        "VSS é ainda pertencente à rede mundial de observatórios magnéticos INTERMAGNET (www.intermagnet.org) e seus instrumentos são rastreados, às 24 horas do dia, ao minuto."
        )
        caixa_texto13 = tk.Text(labelframe10, width=68, height=8, font=('Arial', 12))
        caixa_texto13.grid(row=0, column=1, padx=3, pady=4, sticky='w')
        caixa_texto13.insert(tk.END, texto_inicial)  # Insere o texto inicial
        
        caixa_texto14 = tk.Text(labelframe11, width=68, height=8, font=('Arial', 12))
        caixa_texto14.grid(row=0, column=1, padx=3, pady=4, sticky='w')
        
        caixa_texto15 = tk.Text(labelframe12, width=68, height=8, font=('Arial', 12))
        caixa_texto15.grid(row=0, column=1, padx=3, pady=4, sticky='w')
        
        global tab2
        tab2 = Frame(tabsystem,width=995,height=800,bg='white')
        tabsystem.add(tab2, text='  Derivação  ')
        tabsystem.place(x=5,y=12)
        Frame(tab2, width=642, height=550, background="#dce6f4", highlightbackground="black", highlightthickness=3).place(x=3, y=3)
        
        
        tab4 = Frame(tabsystem,width=995,height=800)
        tabsystem.add(tab4, text='  Observações')
        tabsystem.place(x=5,y=12)
        Frame(tab4, width=642, height=550, background="#dce6f4", highlightbackground="black", highlightthickness=3).place(x=3, y=3)
        
        labelframe13 = LabelFrame(tab4, text=" Observações ",bg='#dce6f4',font=my_font2)
        labelframe13.place(x=10, y=100)
        
        caixa_texto16 = tk.Text(labelframe13, width=68, height=20, font=('Arial', 12))
        caixa_texto16.grid(row=0, column=1, padx=3, pady=4, sticky='w')
        
        Label(tab4,text = "RECORD OF SURVEY CHECK",bg='#dce6f4',font = my_font2).place(x=10, y=10)
        
        labelframe14 = LabelFrame(tab4, text=" Name of last adjust ",bg='#dce6f4',font=my_font2)
        labelframe14.place(x=10, y=130)
        
        
        caixa_texto14 = ttk.Entry(labelframe14, width=14, font=('Arial', 12))
        caixa_texto14.place(x=10, y=160)
    

        
        #################### grafico #############################################
        head = list(range(0, 346, 15))
        DEVIATION = [2.1,2, 1.8, 1.4, 0.8, 0.2, -0.5,-1, -1.3, -1.4, -1.2, -0.8,-0.3
                         ,0.4,0.9,1.4,1.7,1.8,1.9,1.9,1.9,1.9,2,2]
        
        def load_data():
            # Criar DataFrame com base em head e DEVIATION
            data = {'HEAD': head, 'DEVIATION': DEVIATION}
            df = pd.DataFrame(data)
            return df

        frame = Frame(tab2)
        frame.pack(fill='both', expand=True)

        pt = Table(frame, width=550, height=550, showtoolbar=True, showstatusbar=True)
        pt.show()

        data = load_data()
        pt.model.df = data
        
        def plot_grafico():
            # Obter os dados da tabela
            data_table = pt.model.df

            # Limpar o gráfico
            subplot.clear()
            subplot.set_title("Magnetic Compass Deviation Card")
            # Atualizar o gráfico com os dados da tabela
            subplot.scatter(data_table['DEVIATION'],data_table['HEAD'], color='red')
            subplot.plot(data_table['DEVIATION'],data_table['HEAD'], color='black', linewidth=1)
            
            subplot.set_ylim(subplot.get_ylim()[::-1])
            subplot.xaxis.tick_top()
            subplot.xaxis.set_label_position('top')
            # Definir os limites do eixo z
            subplot.set_xlim([-3, 3])
            subplot.grid(True)

            # Redesenhar o canvas
            canvas.draw()

        def inicializar_grafico():
            global figura, subplot, canvas

            x = []
            y = []

            figura = Figure(figsize=(4.5, 5.2))
            subplot = figura.add_subplot(1, 1, 1)
            subplot.plot(x, y)
            subplot.set_title("Magnetic Compass Deviation Card")
            subplot.set_ylim(subplot.get_ylim()[::-1])
            subplot.xaxis.tick_top()
            subplot.xaxis.set_label_position('top')

            canvas = FigureCanvasTkAgg(figura, master=tab2)
            canvas.draw()
            canvas.get_tk_widget().place(x=216, y=0)

      
        # Adicionar botão de atualização do gráfico
        botao_atualizar = tk.Button(tab2, text="Atualizar Gráfico", command=plot_grafico)
        botao_atualizar.place(x=400, y=520)

        # Inicializar a área do gráfico
        inicializar_grafico()


        
        

    frame1 = Frame(tab1, width=642, height=450, background="#dce6f4", highlightbackground="black", highlightthickness=3).place(x=3, y=60)
    
   
    def click1():
        webbrowser.open_new(r"https://github.com/LDSM-ON/Geo-Resistivity-meter")
    def click2():
        webbrowser.open_new(r"https://www.gov.br/observatorio/pt-br")
            
       
        

    labelframe = LabelFrame(tab1, text=" LABORATÓRIOS DE MEDIÇÃO ",bg='#dce6f4',font=my_font2)
    labelframe.place(x=10, y=70)
    labelframe2 = LabelFrame(tab1, text=" FABRICANTE ",bg='#dce6f4',font=my_font2)
    labelframe2.place(x=10, y=150)
    labelframe3 = LabelFrame(tab1, text=" MODELO (MPN) ",bg='#dce6f4',font=my_font2)
    labelframe3.place(x=155, y=150)
    labelframe4 = LabelFrame(tab1, text=" N/S Fab ",bg='#dce6f4',font=my_font2)
    labelframe4.place(x=305, y=150)
    labelframe5 = LabelFrame(tab1, text=" PRÓX. CALIBRAÇÃO ",bg='#dce6f4',font=my_font2)
    labelframe5.place(x=453, y=150)
    labelframe6 = LabelFrame(tab1, text=" REQUERENTE ",bg='#dce6f4',font=my_font2)
    labelframe6.place(x=10, y=215)
    labelframe7 = LabelFrame(tab1, text=" EXECUÇÃO ",bg='#dce6f4',font=my_font2)
    labelframe7.place(x=10, y=280)
    labelframe8 = LabelFrame(tab1, text=" RESPONSÁVEL ",bg='#dce6f4',font=my_font2)
    labelframe8.place(x=10, y=375)
    labelframe9 = LabelFrame(tab1, text=" CERT. DE CAL. ",bg='#dce6f4',font=my_font2)
    labelframe9.place(x=495, y=215)

    

    

        
    def selectPath():   
        path_ = askdirectory()
        path.set(path_)
            
    def create_file():
        dirs = os.path.join(path.get(), folder.get())
        if not os.path.exists(dirs):
            os.makedirs(dirs)
            tkinter.messagebox.showinfo('Tips:','Pasta criado com sucesso!')
            
            
            
    path = StringVar()   # Receiving user's file_path selection
    folder = StringVar() # Receiving user's folder_name selection
    numberChosen2 = IntVar() # Numero de eletrodos
    AB = StringVar()
    na = StringVar()


    
    Label(tab1,text = "Bússola-Meter",font=my_font1).place(x=5, y=5)
    

    def salvar_texto1():
        texto = caixa_texto.get("1.0", "end-1c")  # Obtém todo o texto da caixa de texto
        texto_var.set(texto)

        with open("texto_salvo.txt", "w") as arquivo:
            arquivo.write(texto)

        print(f"Texto salvo: {texto}")

    def carregar_texto1():
        global texto
        try:
            with open("texto_salvo.txt", "r") as arquivo:
                texto = arquivo.read()
                texto_var.set(texto)
                caixa_texto.delete("1.0", "end")  # Limpa o conteúdo anterior
                caixa_texto.insert("1.0", texto)  # Insere o novo texto
                print(f"Texto carregado: {texto}")
        except FileNotFoundError:
            print("Arquivo não encontrado. Não há texto salvo anteriormente.")

    
    def salvar_texto2():
        # Obtém o texto da caixa de texto e salva na variável
        texto2 = caixa_texto2.get()
        texto_var2.set(texto2)

        # Salvar o texto em um arquivo
        with open("texto_salvo2.txt", "w") as arquivo:
            arquivo.write(texto2)

        print(f"Texto salvo: {texto2}")

    def carregar_texto2():
        global texto2
        # Tenta carregar o texto salvo de um arquivo
        try:
            with open("texto_salvo2.txt", "r") as arquivo:
                texto2 = arquivo.read()
                texto_var2.set(texto2)
                print(f"Texto carregado: {texto2}")
        except FileNotFoundError:
            print("Arquivo não encontrado. Não há texto salvo anteriormente.")

    def salvar_texto3():
        # Obtém o texto da caixa de texto e salva na variável
        texto3 = caixa_texto3.get()
        texto_var3.set(texto3)

        # Salvar o texto em um arquivo
        with open("texto_salvo3.txt", "w") as arquivo:
            arquivo.write(texto3)

        print(f"Texto salvo: {texto3}")

    def carregar_texto3():
        global texto3
        # Tenta carregar o texto salvo de um arquivo
        try:
            with open("texto_salvo3.txt", "r") as arquivo:
                texto3 = arquivo.read()
                texto_var3.set(texto3)
                print(f"Texto carregado: {texto3}")
        except FileNotFoundError:
            print("Arquivo não encontrado. Não há texto salvo anteriormente.")

    def salvar_texto4():
        # Obtém o texto da caixa de texto e salva na variável
        texto4 = caixa_texto4.get()
        texto_var4.set(texto4)

        # Salvar o texto em um arquivo
        with open("texto_salvo4.txt", "w") as arquivo:
            arquivo.write(texto4)

        print(f"Texto salvo: {texto4}")

    def carregar_texto4():
        global texto4
        # Tenta carregar o texto salvo de um arquivo
        try:
            with open("texto_salvo4.txt", "r") as arquivo:
                texto4 = arquivo.read()
                texto_var4.set(texto4)
                print(f"Texto carregado: {texto4}")
        except FileNotFoundError:
            print("Arquivo não encontrado. Não há texto salvo anteriormente_4.")
    
    def salvar_texto5():
        # Obtém o texto da caixa de texto e salva na variável
        texto5 = caixa_texto5.get()
        texto_var5.set(texto5)

        # Salvar o texto em um arquivo
        with open("texto_salvo5.txt", "w") as arquivo:
            arquivo.write(texto5)

        print(f"Texto salvo: {texto5}")

    def carregar_texto5():
        global texto5
        # Tenta carregar o texto salvo de um arquivo
        try:
            with open("texto_salvo5.txt", "r") as arquivo:
                texto5 = arquivo.read()
                texto_var5.set(texto5)
                print(f"Texto carregado: {texto5}")
        except FileNotFoundError:
            print("Arquivo não encontrado. Não há texto salvo anteriormente_5.")
            
    def salvar_texto6():
        # Obtém o texto da caixa de texto e salva na variável
        texto6 = caixa_texto6.get()
        texto_var6.set(texto6)

        # Salvar o texto em um arquivo
        with open("texto_salvo6.txt", "w") as arquivo:
            arquivo.write(texto6)

        print(f"Texto salvo: {texto6}")

    def carregar_texto6():
        global texto6
        # Tenta carregar o texto salvo de um arquivo
        try:
            with open("texto_salvo6.txt", "r") as arquivo:
                texto6 = arquivo.read()
                texto_var6.set(texto6)
                print(f"Texto carregado: {texto6}")
        except FileNotFoundError:
            print("Arquivo não encontrado. Não há texto salvo anteriormente_6.")
            
    def salvar_texto7():
        # Obtém o texto da caixa de texto e salva na variável
        texto7 = caixa_texto7.get()
        texto_var7.set(texto7)

        # Salvar o texto em um arquivo
        with open("texto_salvo7.txt", "w") as arquivo:
            arquivo.write(texto7)

        print(f"Texto salvo: {texto7}")

    def carregar_texto7():
        global texto7
        # Tenta carregar o texto salvo de um arquivo
        try:
            with open("texto_salvo7.txt", "r") as arquivo:
                texto7 = arquivo.read()
                texto_var7.set(texto7)
                print(f"Texto carregado: {texto7}")
        except FileNotFoundError:
            print("Arquivo não encontrado. Não há texto salvo anteriormente_7.")        
    
    def salvar_texto8():
        # Obtém o texto da caixa de texto e salva na variável
        texto8 = caixa_texto8.get()
        texto_var8.set(texto8)

        # Salvar o texto em um arquivo
        with open("texto_salvo8.txt", "w") as arquivo:
            arquivo.write(texto8)

        print(f"Texto salvo: {texto8}")

    def carregar_texto8():
        global texto8
        # Tenta carregar o texto salvo de um arquivo
        try:
            with open("texto_salvo8.txt", "r") as arquivo:
                texto8 = arquivo.read()
                texto_var8.set(texto8)
                print(f"Texto carregado: {texto8}")
        except FileNotFoundError:
            print("Arquivo não encontrado. Não há texto salvo anteriormente_8.")
            
    def salvar_texto9():
        # Obtém o texto da caixa de texto e salva na variável
        texto9 = caixa_texto9.get()
        texto_var9.set(texto9)

        # Salvar o texto em um arquivo
        with open("texto_salvo9.txt", "w") as arquivo:
            arquivo.write(texto9)

        print(f"Texto salvo: {texto9}")

    def carregar_texto9():
        global texto9
        # Tenta carregar o texto salvo de um arquivo
        try:
            with open("texto_salvo9.txt", "r") as arquivo:
                texto9 = arquivo.read()
                texto_var9.set(texto9)
                print(f"Texto carregado: {texto9}")
        except FileNotFoundError:
            print("Arquivo não encontrado. Não há texto salvo anteriormente_9.")
            
    def salvar_texto10():
        # Obtém o texto da caixa de texto e salva na variável
        texto10 = caixa_texto10.get()
        texto_var10.set(texto10)

        # Salvar o texto em um arquivo
        with open("texto_salvo10.txt", "w") as arquivo:
            arquivo.write(texto10)

        print(f"Texto salvo: {texto10}")

    def carregar_texto10():
        global texto10
        # Tenta carregar o texto salvo de um arquivo
        try:
            with open("texto_salvo10.txt", "r") as arquivo:
                texto10 = arquivo.read()
                texto_var10.set(texto10)
                print(f"Texto carregado: {texto10}")
        except FileNotFoundError:
            print("Arquivo não encontrado. Não há texto salvo anteriormente_10.")
            
    def salvar_texto11():
        # Obtém o texto da caixa de texto e salva na variável
        texto11 = caixa_texto11.get()
        texto_var11.set(texto11)

        # Salvar o texto em um arquivo
        with open("texto_salvo11.txt", "w") as arquivo:
            arquivo.write(texto11)

        print(f"Texto salvo: {texto11}")

    def carregar_texto11():
        global texto11
        # Tenta carregar o texto salvo de um arquivo
        try:
            with open("texto_salvo11.txt", "r") as arquivo:
                texto11 = arquivo.read()
                texto_var11.set(texto11)
                print(f"Texto carregado: {texto11}")
        except FileNotFoundError:
            print("Arquivo não encontrado. Não há texto salvo anteriormente_11.")
            
    def salvar_texto12():
        # Obtém o texto da caixa de texto e salva na variável
        texto12 = caixa_texto12.get()
        texto_var12.set(texto12)

        # Salvar o texto em um arquivo
        with open("texto_salvo12.txt", "w") as arquivo:
            arquivo.write(texto12)

        print(f"Texto salvo: {texto12}")

    def carregar_texto12():
        global texto12
        # Tenta carregar o texto salvo de um arquivo
        try:
            with open("texto_salvo12.txt", "r") as arquivo:
                texto12 = arquivo.read()
                texto_var12.set(texto12)
                print(f"Texto carregado: {texto12}")
        except FileNotFoundError:
            print("Arquivo não encontrado. Não há texto salvo anteriormente_12.")          
    
    def salvar_texto13():
        # Obtém o texto da caixa de texto e salva na variável
        global texto13
        texto13 = caixa_texto13.get()
        texto_var13.set(texto13)

        # Salvar o texto em um arquivo
        with open("texto_salvo13.txt", "w") as arquivo:
            arquivo.write(texto13)

        print(f"Texto salvo: {texto13}")

    def carregar_texto13():
        global texto13
        # Tenta carregar o texto salvo de um arquivo
        try:
            with open("texto_salvo13.txt", "r") as arquivo:
                texto13 = arquivo.read()
                texto_var13.set(texto13)
                print(f"Texto carregado: {texto13}")
        except FileNotFoundError:
            print("Arquivo não encontrado. Não há texto salvo anteriormente_13.")
            
            
    texto_var = tk.StringVar()
    caixa_texto = tk.Text(labelframe, wrap="word", width=68, height=2, font=('Arial', 12))
    caixa_texto.grid(row=1, column=0, padx=3, pady=4, sticky='w')
        
    texto_var2 = tk.StringVar()
    caixa_texto2 = ttk.Entry(labelframe2, textvariable=texto_var2, width=14, font=('Arial', 12))
    caixa_texto2.grid(row=3, column=0, padx=3, pady=4, sticky='w')
    
    texto_var3 = tk.StringVar()
    caixa_texto3 = ttk.Entry(labelframe3, textvariable=texto_var3, width=14, font=('Arial', 12))
    caixa_texto3.grid(row=5, column=0, padx=3, pady=4, sticky='w')
    
    texto_var4 = tk.StringVar()
    caixa_texto4 = ttk.Entry(labelframe4, textvariable=texto_var4, width=14, font=('Arial', 12))
    caixa_texto4.grid(row=7, column=0, padx=3, pady=4, sticky='w')
       
    texto_var5 = tk.StringVar()
    caixa_texto5 = ttk.Entry(labelframe5, textvariable=texto_var5, width=14, font=('Arial', 12))
    caixa_texto5.grid(row=7, column=0, padx=3, pady=4, sticky='w')
    
    texto_var6 = tk.StringVar()
    caixa_texto6 = ttk.Entry(labelframe6, textvariable=texto_var6, width=51, font=('Arial', 12))
    caixa_texto6.grid(row=7, column=0, padx=3, pady=4, sticky='w')
    
    texto_var7 = tk.StringVar()
    caixa_texto7 = ttk.Entry(labelframe7, textvariable=texto_var7, width=42, font=('Arial', 12))
    caixa_texto7.grid(row=0, column=0, padx=3, pady=4, sticky='w')
    
    texto_var8 = tk.StringVar()
    caixa_texto8 = ttk.Entry(labelframe7, textvariable=texto_var8, width=25, font=('Arial', 12))
    caixa_texto8.grid(row=0, column=1, padx=3, pady=4, sticky='w')
    
    texto_var9 = tk.StringVar()
    caixa_texto9 = ttk.Entry(labelframe7, textvariable=texto_var9, width=42, font=('Arial', 12))
    caixa_texto9.grid(row=1, column=0, padx=3, pady=4, sticky='w')
    
    texto_var10 = tk.StringVar()
    caixa_texto10 = ttk.Entry(labelframe7, textvariable=texto_var10, width=25, font=('Arial', 12))
    caixa_texto10.grid(row=1, column=1, padx=3, pady=4, sticky='w')
    
    texto_var11 = tk.StringVar()
    caixa_texto11 = ttk.Entry(labelframe8, textvariable=texto_var11, width=42, font=('Arial', 12))
    caixa_texto11.grid(row=0, column=0, padx=3, pady=4, sticky='w')
    
    texto_var12 = tk.StringVar()
    caixa_texto12 = ttk.Entry(labelframe8, textvariable=texto_var12, width=25, font=('Arial', 12))
    caixa_texto12.grid(row=0, column=1, padx=3, pady=4, sticky='w')
    
    texto_var13 = tk.StringVar()
    caixa_texto13 = ttk.Entry(labelframe9, textvariable=texto_var13, width=14, font=('Arial', 12))
    caixa_texto13.grid(row=0, column=1, padx=3, pady=4, sticky='w')
    
    
    

    


    carregar_texto1()
    carregar_texto2()
    carregar_texto3()
    carregar_texto4()
    carregar_texto5()
    carregar_texto6()
    carregar_texto7()
    carregar_texto8()
    carregar_texto9()
    carregar_texto10()
    carregar_texto11()
    carregar_texto12()
    carregar_texto13()
    
    
  

    botao_salvar = tk.Button(tab1, text="Salvar Texto", command=lambda: [salvar_texto1(),salvar_texto2,salvar_texto3,salvar_texto4(),salvar_texto5()
                                            ,salvar_texto6(),salvar_texto7(),salvar_texto8(),salvar_texto9()
                                            ,salvar_texto10(),salvar_texto11(),salvar_texto12(),salvar_texto13()])
    botao_salvar.place(x=15, y=440, width=100, height=42)


    
    
    ################## TAB 2 #########################################################
  
    
    botao_nova_tab = tk.Button(tab1, text="Deviation Card", command=novaTab)
    botao_nova_tab.place(x=150, y=440, width=100, height=42)
    
    def cria_doc():
        def adicionar_imagem_cabecalho(document, caminho_imagem, largura=Inches(1.25)):
            # Adicionar uma tabela ao cabeçalho
            cabecalho = document.sections[0].header
            tabela = cabecalho.add_table(rows=1, cols=1, width=largura)

            # Obter a célula da tabela
            celula = tabela.cell(0, 0)

            # Adicionar a imagem à célula
            run = celula.paragraphs[0].add_run()
            run.add_picture(caminho_imagem, width=largura)

            # Alinhar a imagem no centro da célula
            paragrafo = celula.paragraphs[0]
            paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Adicionar um espaçamento para separar a imagem do texto (opcional)
            paragrafo.add_run().add_break()

        # Exemplo de uso
        doc = Document()
        
        adicionar_imagem_cabecalho(doc, 'C:\\Users\\ON\\Desktop\\Bússola\\figure\\logo3.png', largura=Inches(6))  # Ajuste o valor de largura confor
        
        paragrafo1 = doc.add_paragraph('CERTIFICADO DE CALIBRAÇÃO  Nº '+str(texto13)+'/'+str(datetime.now().strftime("%Y")))
        paragrafo1.alignment = 1  # 1 representa centralizado (0 é à esquerda, 2 é à direita)
        
        
        # Table data in a form of list 
        data1 = ( 
            (str(texto), '1 de 1'), 
            
        ) 

        # Creating the first table 
        table1 = doc.add_table(rows=1, cols=2) 

        # Adding heading in the 1st row of the first table 
        header_row1 = table1.rows[0].cells 
        header_row1[0].text = 'LABORATÓRIOS DE MEDIÇÃO'
        header_row1[1].text = 'FOLHA'

        header_row1[0].width = docx.shared.Inches(20)

        # Adding data from the list to the first table 
        for id, name in data1: 
            # Adding a row and then adding data in it. 
            row = table1.add_row().cells 
            row[0].text = str(id) 
            row[1].text = name 

        # Adding a new line as a separator
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = 0

        # Table data for the second table 
        data2 = ( 
            (str(texto2), str(texto4)), 
        ) 

        # Creating the second table 
        table2 = doc.add_table(rows=1, cols=2) 

        # Adding heading in the 1st row of the second table 
        header_row2 = table2.rows[0].cells 
        header_row2[0].text = 'FABRICANTE'
        header_row2[1].text = 'N/S Fab'
        header_row2[0].width = docx.shared.Inches(20)




        # Adding data from the list to the second table 
        for id, name in data2: 
            # Adding a row and then adding data in it. 
            row = table2.add_row().cells 
            row[0].text = str(id) 
            row[1].text = name 
        
        # Adding a new line as a separator
        para2 = doc.add_paragraph()
        para2.paragraph_format.line_spacing = 0

        # Table data for the second table 
        data3 = ( 
            (str(texto3), ''), 
        ) 

        # Creating the second table 
        table3 = doc.add_table(rows=1, cols=2) 

        # Adding heading in the 1st row of the second table 
        header_row3 = table3.rows[0].cells 
        header_row3[0].text = 'MODELO (MPN)'
        header_row3[1].text = ''
        header_row3[0].width = docx.shared.Inches(20)




        # Adding data from the list to the second table 
        for id, name in data3: 
            # Adding a row and then adding data in it. 
            row = table3.add_row().cells 
            row[0].text = str(id) 
            row[1].text = name
        
                # Adding a new line as a separator
        para3 = doc.add_paragraph()
        para3.paragraph_format.line_spacing = 0

        # Table data for the second table 
        data4 = ( 
            (str(texto6), str(texto5)), 
        ) 

        # Creating the second table 
        table4 = doc.add_table(rows=1, cols=2) 

        # Adding heading in the 1st row of the second table 
        header_row4 = table4.rows[0].cells 
        header_row4[0].text = 'REQUERENTE'
        header_row4[1].text = 'PRÓX. CAL.'
        header_row4[0].width = docx.shared.Inches(20)




        # Adding data from the list to the second table 
        for id, name in data4: 
            # Adding a row and then adding data in it. 
            row = table4.add_row().cells 
            row[0].text = str(id) 
            row[1].text = name
            
        
        # Adding a new line as a separator
        para4 = doc.add_paragraph()
        para4.paragraph_format.line_spacing = 0

        # Table data for the second table 
        data5 = ( 
            ('ETAPAS', ''),
            ('As medições foram realizadas em duas etapas. \n A primeira no Laboratório de Desenvolvimento de Sensores Magnéticos (LDSM/ON) com instrumentos acreditados. \n A segunda no Observatório Magnético de Vassouras(VSS) em operação desde 1915 e referenciado pela IAGA (International Association of Geomagnetism and Aeronomy). VSS é ainda pertencente à rede mundial de observatórios magnéticos INTERMAGNET (www.intermagnet.org) e seus instrumentos são rastreados, às 24 horas do dia, ao minuto.', ''),
        ('INSTRUMENTOS E LOCAIS', ''),
            ('1 – Bobina Triaxial de Helmholtz  - LDSM/ON \n 2 – Fontes de corrente, de precisão, Keithley – LDSM/ON \n 3 – Magnetômetro DI Flux Bartington, com precisão de um minuto de arco (1’) e pilar com referência astronômica – VSS.', ''),
        ('CALIBRAÇÃO:', ''),
            ('No LDSM/ON o sincronismo de orientação magnética e o “offset” foram testados utilizando uma bobina triaxial de Helmholtz, acreditada, alimentada por fontes de corrente, de precisão, Keithley. \n Em VSS a bússola foi alinhada sob uma base certificada pela IAGA e orientada na direção N-S verdadeira (geográfica). Com auxílio de magnetômetro DI-flux, acreditado e rastreado, determinou-se a qualificação da bússola que mostrou incertezas menores que 2,5º (limitados pela resolução própria do instrumento, de 5º).', ''),
        ) 

        # Creating the second table 
        table5 = doc.add_table(rows=1, cols=2) 

        # Adding heading in the 1st row of the second table 
        header_row5 = table5.rows[0].cells 
        header_row5[0].text = 'PROCEDIMENTOS SEGUIDOS'
        header_row5[1].text = '.'
        header_row5[0].width = docx.shared.Inches(20)




        # Adding data from the list to the second table 
        for id, name in data5: 
            # Adding a row and then adding data in it. 
            row = table5.add_row().cells 
            row[0].text = str(id) 
            row[1].text = name
            
        # Adding style to both tables 
        table1.style = 'Colorful List'
        table2.style = 'Colorful List'
        table3.style = 'Colorful List'
        table4.style = 'Colorful List'
        table5.style = 'Colorful List'


        

        # Adicionar data e hora ao nome do arquivo
        data_hora_atual = datetime.now().strftime("%d-%m-%Y_%H-%M-%S")
        nome_arquivo = f'exemplo_com_imagem_cabecalho_{data_hora_atual}.docx'
        Label(tab1,text = "C:\\Users\\ON\\Desktop\\teste usuario\\"+str(nome_arquivo)).place(x=60, y=520)

        doc.save('C:\\Users\\ON\\Desktop\\teste usuario\\'+str(nome_arquivo))
        
   
    
        
    botao_doc = tk.Button(tab1, text="botao_doc", command=cria_doc)
    botao_doc.place(x=350, y=440, width=100, height=42)
    #>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>
       
     

root = Tk()  # Creates the main window (root) using the Tk() class.
root.title('Bússola-Meter')  # Sets the window title as "Geo-Resistivity-Meter".
width_of_window = 610 # Defines the width and height of the window.
height_of_window = 350
screen_width = root.winfo_screenwidth() # Gets the width and height of the computer screen.
screen_height = root.winfo_screenheight()
x_coordinate = (screen_width/2)-(width_of_window/2)
y_coordinate = (screen_height/2)-(height_of_window/2)
root.geometry("%dx%d+%d+%d" % (width_of_window, height_of_window, x_coordinate, y_coordinate))  # Defines the geometry of the window with the specified width, height, and calculated coordinates.

icon = PhotoImage(file="figure/energia2.png")  # Loads an image "energia2.png" to be used as the window icon.
root.iconphoto(True, icon)  # Sets the loaded image as the window icon.
        
s = Style()  # Creates a Style object, used to configure styles and themes of widgets.
s.theme_use('clam')  # Sets the theme to 'clam', which is one of the available themes in the ttk library.
s.configure("red.Horizontal.TProgressbar", foreground='red', background='#4f4f4f')  # Configures the style for the horizontal progress bar (Progressbar widget) with red foreground color and gray background color.
        
progress = Progressbar(root, style="red.Horizontal.TProgressbar", orient=HORIZONTAL, length=612, mode='determinate')  # Creates a Progressbar widget in the window, using the previously configured style, with horizontal orientation, length of 500 pixels, and determinate mode (indicating a progress value).
progress.place(x=0, y=334)  # Positions the progress bar at the specified (x, y) coordinate in the window.
        
a = '#249794'  # Background color for the frame.
        
frame1 = Frame(root, width=510, height=241)  # Creates a Frame widget in the window with the specified width and height.
frame1.place(x=0, y=0)  # Positions the frame at the specified (x, y) coordinate in the window.
        
home_image = ImageTk.PhotoImage(Image.open("figure/logo5.png"))  # Loads an image "logo2.png" to be used in the Label widget.
Label(frame1, image=home_image).grid(row=1, columnspan=1)  # Creates a Label widget with the loaded image and places it in the second row of the grid of the frame1 widget.
        
def EXT():  # Defines the function to close the window.
    root.destroy()
        
            
def bar():  # Defines the function to update a progress bar while executing other functions.
    label_bar = Label(root, text='Loading...',bg="black",fg="white")
    lst4 = ('Calibri (Body)', 10)
    label_bar.config(font=lst4)
    label_bar.place(x=10, y=300)
    import time
    r = 0
    for i in range(100):
        progress['value'] = r
        root.update_idletasks()
        time.sleep(0.03)
        r += 1
            
    EXT()
    run()
b1 = Button(root, text='Iniciar', command=bar)  # Creates a button in the window with the text "Iniciar" and assigns the bar() function to be executed when clicked.
b1.place(x=10, y=240, width=100, height=42)  # Positions the button at the specified (x, y) coordinate with the specified width and height.
        
      

root.mainloop()


    # Iniciar o loop de eventos do Tkinter
    

